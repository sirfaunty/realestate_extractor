"""
Generate the Lease Abstract Compendium (Word .docx) from the warehouse.

Mirrors the partner's deliverable #6: a title block, a "How to Read" legend, then
every provision grouped by Article — each provision shown as a heading followed by a
three-row table (Detailed / Detailed Summary / Abstract Summary).

Reads the local warehouse built by build_warehouse.py + abstract_lease.py. Works on
either the locally-generated abstracts (engine = llama3.1:8b) or the gold reference
warehouse (which has no engine column).

IMPORTANT: abstracts here are machine-generated (local model). The document is labeled
an AUTOMATED FIRST DRAFT FOR INTERNAL REVIEW — not attorney work product.

Usage:
    python compendium_docx.py --db data/lease_warehouse.db --out Southtown_Lease_Compendium.docx
    python compendium_docx.py --db data/gold_lease_warehouse.db --engine gold --out gold_compendium.docx
"""
import argparse
import datetime
import os
import sqlite3

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

KA_NAVY = RGBColor(0x1B, 0x2A, 0x4A)
KA_RED = RGBColor(0xC8, 0x10, 0x2E)
KA_BLUE = RGBColor(0x2E, 0x5A, 0x88)
NAVY_FILL = "1B2A4A"
CLOUD_FILL = "EEF3F8"

TIERS = [("detailed", "Detailed"),
         ("detailed_summary", "Detailed Summary"),
         ("abstract_summary", "Abstract Summary")]

LEGEND = {
    "Detailed": "Nearly all the meaningful lease language — operative terms, defined "
                "terms, thresholds, deadlines, and each party option/remedy — with the "
                "section cited.",
    "Detailed Summary": "The substance with legal jargon stripped out: what the "
                        "provision does, its triggers, and why it matters.",
    "Abstract Summary": "One-line summary for matrices, briefing decks, or summary "
                        "reports.",
}


def _has_engine_col(conn):
    return any(r[1] == "engine" for r in conn.execute("PRAGMA table_info(abstracts)"))


def _load(conn, engine):
    """Return ordered provisions, each with {tier: content}."""
    provs = conn.execute(
        "SELECT id, article_num, article_roman, article_title, section_num, "
        "section_heading, body FROM provisions ORDER BY seq").fetchall()
    use_engine = _has_engine_col(conn) and engine and engine != "gold"
    out = []
    for pid, anum, aroman, atitle, snum, head, body in provs:
        if use_engine:
            rows = conn.execute(
                "SELECT abstract_type, content FROM abstracts WHERE provision_id=? AND engine=?",
                (pid, engine)).fetchall()
        else:
            rows = conn.execute(
                "SELECT abstract_type, content FROM abstracts WHERE provision_id=?",
                (pid,)).fetchall()
        out.append({
            "article_num": anum, "article_roman": aroman, "article_title": atitle,
            "section_num": snum, "section_heading": head or "", "body": body or "",
            "abstracts": {t: c for t, c in rows},
        })
    return out


def _shade(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcpr.append(shd)


def _page_number_footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CONFIDENTIAL — Automated draft for internal review     ·     Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x60, 0x6A, 0x78)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)


def _title_block(doc, title, subtitle_lines):
    def line(text, size, color, bold=True, after=2, upper=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(after)
        r = p.add_run(text.upper() if upper else text)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.color.rgb = color
        return p

    line("KRAUS-ANDERSON", 22, KA_NAVY)
    line("Portfolio Development  |  Southtown Shopping Center", 11, KA_BLUE, bold=False, after=10)
    line(title, 18, KA_NAVY, after=2)
    for s in subtitle_lines:
        line(s, 11, KA_BLUE, bold=False, after=2)
    # red rule
    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr = rule.paragraph_format
    pr.space_before = Pt(8)
    pborders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "C8102E")
    pborders.append(bottom)
    rule._p.get_or_add_pPr().append(pborders)
    banner = line("AUTOMATED FIRST DRAFT FOR INTERNAL REVIEW — NOT ATTORNEY WORK PRODUCT",
                  9, KA_RED, after=2)
    banner.paragraph_format.space_before = Pt(8)


def _provision_table(doc, prov):
    tbl = doc.add_table(rows=0, cols=2)
    tbl.autofit = False
    for key, label in TIERS:
        row = tbl.add_row().cells
        row[0].width = Inches(1.5)
        row[1].width = Inches(5.0)
        _shade(row[0], NAVY_FILL)
        lp = row[0].paragraphs[0]
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(9)
        lr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cp = row[1].paragraphs[0]
        cr = cp.add_run(prov["abstracts"].get(key) or "—")
        cr.font.size = Pt(10)
    # thin borders
    tbl.style = "Table Grid"
    return tbl


def _is_deleted(prov):
    h = prov["section_heading"].lower()
    return "intentionally" in h or (len(prov["body"].strip()) < 15 and "deleted" in h)


def build(db_path, out_path, engine="llama3.1:8b"):
    conn = sqlite3.connect(db_path)
    provs = _load(conn, engine)
    conn.close()

    doc = Document()
    # base font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    _title_block(doc, "Dick's House of Sport Lease — Provision-by-Provision Abstract Compendium",
                 ["Three-tier abstracts — Detailed | Detailed Summary | Abstract Summary",
                  f"All {len(provs)} provisions, Articles I–XVII — abstracted from the lease text",
                  f"Generated {datetime.date.today().isoformat()} · engine: {engine}"])

    doc.add_heading("How to Read This Compendium", level=1)
    intro = doc.add_paragraph()
    intro.add_run(
        "Each provision is abstracted at three levels of detail. Abstracts are generated "
        "locally by a language model directly from the lease text and are a first draft "
        "for human review — verify every figure and defined term against the source "
        "before relying on it. Provisions marked “Intentionally Deleted” in the "
        "lease are retained for completeness.").font.size = Pt(10)
    legend = doc.add_table(rows=0, cols=2)
    legend.style = "Table Grid"
    for label, desc in LEGEND.items():
        c = legend.add_row().cells
        c[0].width = Inches(1.5); c[1].width = Inches(5.0)
        _shade(c[0], CLOUD_FILL)
        r0 = c[0].paragraphs[0].add_run(label); r0.bold = True; r0.font.size = Pt(9)
        c[1].paragraphs[0].add_run(desc).font.size = Pt(9)

    # Group by article, preserving order
    cur_article = None
    for prov in provs:
        anum = prov["article_num"]
        if anum != cur_article:
            cur_article = anum
            if anum == 0:
                title = "Preamble / Recitals"
            else:
                title = f"Article {prov['article_roman']} — {(prov['article_title'] or '').upper()}"
            doc.add_page_break()
            doc.add_heading(title, level=1)

        head = prov["section_heading"]
        sec = "" if prov["section_num"] in ("0.0",) else f"§{prov['section_num']}  "
        h2 = doc.add_heading(f"{sec}{head}", level=2)
        for run in h2.runs:
            run.font.color.rgb = KA_NAVY
        if _is_deleted(prov):
            note = doc.add_paragraph()
            nr = note.add_run("Intentionally Omitted / Deleted.")
            nr.italic = True
            nr.font.color.rgb = RGBColor(0x60, 0x6A, 0x78)
        else:
            _provision_table(doc, prov)

    _page_number_footer(doc.sections[0])
    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)
    return {"out": out_path, "provisions": len(provs)}


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--db", required=True)
    ap.add_argument("--engine", default="llama3.1:8b",
                    help="Abstract engine to render (or 'gold' for the reference DB)")
    ap.add_argument("--out", default="Southtown_Lease_Abstract_Compendium.docx")
    args = ap.parse_args()
    rep = build(args.db, args.out, engine=args.engine)
    print(f"Wrote {rep['out']} ({rep['provisions']} provisions)")
