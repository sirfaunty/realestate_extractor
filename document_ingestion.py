"""
Multi-format document ingestion.

Supports: PDF, XLSX, DOCX, DOC, MSG, PNG/JPG, MD/TXT/CSV

Each ingestion function returns a DocumentContent object so the rest
of the pipeline (storage, classification, analysis) works unchanged.
"""

import os
import hashlib
import logging
from typing import Optional

from .pdf_ingestion import (
    DocumentContent, PageContent, compute_file_hash, ingest_pdf
)

logger = logging.getLogger(__name__)

# Map extensions to ingestion functions
SUPPORTED_EXTENSIONS = {
    'pdf', 'xlsx', 'xls', 'docx', 'doc',
    'msg', 'png', 'jpg', 'jpeg',
    'md', 'txt', 'csv', 'tsv', 'rtf',
}


def ingest_document(filepath: str, force_ocr: bool = False,
                    preprocess: bool = True) -> DocumentContent:
    """
    Main entry point — routes to the correct ingestion function
    based on file extension.

    Returns a DocumentContent with text (and tables where applicable).
    """
    ext = os.path.splitext(filepath)[1].lower().lstrip('.')

    if ext == 'pdf':
        return ingest_pdf(filepath, force_ocr=force_ocr, preprocess=preprocess)
    elif ext in ('xlsx', 'xls'):
        return ingest_xlsx(filepath)
    elif ext == 'docx':
        return ingest_docx(filepath)
    elif ext == 'doc':
        return ingest_doc(filepath)
    elif ext == 'msg':
        return ingest_msg(filepath)
    elif ext in ('png', 'jpg', 'jpeg'):
        return ingest_image(filepath)
    elif ext in ('md', 'txt', 'csv', 'tsv', 'rtf'):
        return ingest_text(filepath)
    else:
        raise ValueError(f"Unsupported file type: .{ext}")


# ─── XLSX / XLS ─────────────────────────────────────────────────────

def ingest_xlsx(filepath: str) -> DocumentContent:
    """
    Extract text + tables from Excel workbooks.

    Each sheet becomes a "page". Cell values are read as text.
    Sheets with tabular data are also stored as tables.
    """
    import openpyxl

    file_hash = compute_file_hash(filepath)
    filename = os.path.basename(filepath)

    try:
        wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
    except Exception as e:
        logger.warning(f"openpyxl failed on {filename}, trying xlrd: {e}")
        return _ingest_xls_fallback(filepath, file_hash)

    pages = []
    for sheet_idx, sheet_name in enumerate(wb.sheetnames):
        ws = wb[sheet_name]
        rows_data = []
        text_lines = [f"Sheet: {sheet_name}", ""]

        try:
            for row in ws.iter_rows(values_only=True):
                # Convert each cell to string
                str_row = []
                for cell in row:
                    if cell is None:
                        str_row.append('')
                    elif isinstance(cell, float):
                        # Avoid floating point noise
                        if cell == int(cell):
                            str_row.append(str(int(cell)))
                        else:
                            str_row.append(f"{cell:,.2f}")
                    else:
                        str_row.append(str(cell))
                rows_data.append(str_row)

                # Also build text representation
                non_empty = [s for s in str_row if s.strip()]
                if non_empty:
                    text_lines.append('  |  '.join(non_empty))
        except Exception as e:
            logger.warning(f"Error reading sheet '{sheet_name}': {e}")
            text_lines.append(f"[Error reading sheet: {e}]")

        # Build table (first row = headers, rest = data)
        tables = []
        if rows_data and len(rows_data) >= 2:
            tables.append(rows_data)

        pages.append(PageContent(
            page_number=sheet_idx + 1,
            text='\n'.join(text_lines),
            tables=tables,
            is_scanned=False,
        ))

    wb.close()

    return DocumentContent(
        filepath=filepath,
        filename=filename,
        pages=pages,
        page_count=len(pages),
        is_scanned=False,
        file_hash=file_hash,
    )


def _ingest_xls_fallback(filepath: str, file_hash: str) -> DocumentContent:
    """Fallback for old .xls files using xlrd (if available)."""
    filename = os.path.basename(filepath)
    try:
        import xlrd
        wb = xlrd.open_workbook(filepath)
        pages = []
        for sheet_idx in range(wb.nsheets):
            ws = wb.sheet_by_index(sheet_idx)
            text_lines = [f"Sheet: {ws.name}", ""]
            rows_data = []
            for r in range(ws.nrows):
                row_vals = [str(ws.cell_value(r, c)) for c in range(ws.ncols)]
                rows_data.append(row_vals)
                non_empty = [s for s in row_vals if s.strip()]
                if non_empty:
                    text_lines.append('  |  '.join(non_empty))
            tables = [rows_data] if len(rows_data) >= 2 else []
            pages.append(PageContent(
                page_number=sheet_idx + 1,
                text='\n'.join(text_lines),
                tables=tables,
                is_scanned=False,
            ))
        return DocumentContent(
            filepath=filepath, filename=filename,
            pages=pages, page_count=len(pages),
            is_scanned=False, file_hash=file_hash,
        )
    except ImportError:
        logger.warning("xlrd not installed — cannot read .xls files")
        return DocumentContent(
            filepath=filepath, filename=filename,
            pages=[PageContent(page_number=1, text=f"[.xls file — xlrd not installed]")],
            page_count=1, is_scanned=False, file_hash=file_hash,
        )


# ─── DOCX ───────────────────────────────────────────────────────────

def ingest_docx(filepath: str) -> DocumentContent:
    """
    Extract text + tables from Word documents.

    All body text goes into page 1 (Word doesn't have fixed pages).
    Each table in the document is captured separately.
    """
    from docx import Document as DocxDocument

    file_hash = compute_file_hash(filepath)
    filename = os.path.basename(filepath)

    doc = DocxDocument(filepath)

    # Extract body text
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    full_text = '\n'.join(paragraphs)

    # Extract tables
    tables = []
    for tbl in doc.tables:
        rows = []
        for row in tbl.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            tables.append(rows)

    # Treat the whole document as one page
    pages = [PageContent(
        page_number=1,
        text=full_text,
        tables=tables,
        is_scanned=False,
    )]

    return DocumentContent(
        filepath=filepath,
        filename=filename,
        pages=pages,
        page_count=1,
        is_scanned=False,
        file_hash=file_hash,
    )


# ─── DOC (legacy Word) ─────────────────────────────────────────────

def ingest_doc(filepath: str) -> DocumentContent:
    """
    Extract text from legacy .doc files.

    Tries antiword first, falls back to textract, then to raw text extraction.
    Tables are not reliably extractable from .doc files.
    """
    import subprocess

    file_hash = compute_file_hash(filepath)
    filename = os.path.basename(filepath)
    text = ''

    # Try antiword
    try:
        result = subprocess.run(
            ['antiword', filepath],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0 and result.stdout.strip():
            text = result.stdout.strip()
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    # Fallback: try catdoc
    if not text:
        try:
            result = subprocess.run(
                ['catdoc', filepath],
                capture_output=True, text=True, timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                text = result.stdout.strip()
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

    # Fallback: try python-docx on .doc (sometimes works if it's actually docx)
    if not text:
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(filepath)
            text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception:
            pass

    # Last resort: raw binary text extraction
    if not text:
        try:
            with open(filepath, 'rb') as f:
                raw = f.read()
            # Extract printable ASCII runs
            import re
            text_chunks = re.findall(rb'[\x20-\x7e]{20,}', raw)
            text = '\n'.join(chunk.decode('ascii', errors='ignore') for chunk in text_chunks)
            if text:
                text = f"[Extracted from binary — formatting lost]\n\n{text}"
            else:
                text = f"[.doc file — no text extraction tool available. Install antiword: brew install antiword]"
        except Exception as e:
            text = f"[.doc file — extraction failed: {e}]"

    pages = [PageContent(
        page_number=1,
        text=text,
        is_scanned=False,
    )]

    return DocumentContent(
        filepath=filepath,
        filename=filename,
        pages=pages,
        page_count=1,
        is_scanned=False,
        file_hash=file_hash,
    )


# ─── MSG (Outlook email) ────────────────────────────────────────────

def ingest_msg(filepath: str) -> DocumentContent:
    """
    Extract text from Outlook .msg email files.

    Extracts sender, recipients, subject, date, and body text.
    Attachments are listed but not recursively extracted.
    """
    import extract_msg

    file_hash = compute_file_hash(filepath)
    filename = os.path.basename(filepath)

    try:
        msg = extract_msg.Message(filepath)

        # Build structured text from email fields
        lines = []
        if msg.subject:
            lines.append(f"Subject: {msg.subject}")
        if msg.sender:
            lines.append(f"From: {msg.sender}")
        if msg.to:
            lines.append(f"To: {msg.to}")
        if msg.cc:
            lines.append(f"CC: {msg.cc}")
        if msg.date:
            lines.append(f"Date: {msg.date}")

        lines.append("")
        lines.append("--- Body ---")
        lines.append("")

        body = msg.body or ''
        if body:
            lines.append(body.strip())
        else:
            # Try HTML body
            html_body = msg.htmlBody
            if html_body:
                # Strip HTML tags for plain text
                import re
                clean = re.sub(r'<[^>]+>', ' ', html_body if isinstance(html_body, str) else html_body.decode('utf-8', errors='ignore'))
                clean = re.sub(r'\s+', ' ', clean).strip()
                lines.append(clean)
            else:
                lines.append("[No body text found]")

        # List attachments
        if msg.attachments:
            lines.append("")
            lines.append("--- Attachments ---")
            for att in msg.attachments:
                att_name = getattr(att, 'longFilename', None) or getattr(att, 'shortFilename', None) or 'unnamed'
                lines.append(f"  - {att_name}")

        text = '\n'.join(lines)
        msg.close()

    except Exception as e:
        logger.warning(f"Failed to parse MSG {filename}: {e}")
        text = f"[MSG parsing failed: {e}]"

    pages = [PageContent(
        page_number=1,
        text=text,
        is_scanned=False,
    )]

    return DocumentContent(
        filepath=filepath,
        filename=filename,
        pages=pages,
        page_count=1,
        is_scanned=False,
        file_hash=file_hash,
    )


# ─── Images (PNG, JPG) ──────────────────────────────────────────────

def ingest_image(filepath: str) -> DocumentContent:
    """
    Extract text from images using OCR (Tesseract).

    Uses the same OCR pipeline as scanned PDFs.
    """
    file_hash = compute_file_hash(filepath)
    filename = os.path.basename(filepath)

    text = ''
    ocr_confidence = None

    try:
        import pytesseract
        from PIL import Image

        img = Image.open(filepath)

        # Get OCR data with confidence
        data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
        words = []
        confidences = []
        for i, word in enumerate(data['text']):
            if word.strip():
                words.append(word)
                conf = data['conf'][i]
                if isinstance(conf, (int, float)) and conf >= 0:
                    confidences.append(conf)

        text = ' '.join(words)
        if confidences:
            ocr_confidence = sum(confidences) / len(confidences)

    except ImportError:
        text = "[Image file — pytesseract not installed for OCR]"
    except Exception as e:
        logger.warning(f"OCR failed on {filename}: {e}")
        text = f"[Image OCR failed: {e}]"

    pages = [PageContent(
        page_number=1,
        text=text,
        is_scanned=True,
        ocr_confidence=ocr_confidence,
    )]

    return DocumentContent(
        filepath=filepath,
        filename=filename,
        pages=pages,
        page_count=1,
        is_scanned=True,
        avg_ocr_confidence=ocr_confidence,
        file_hash=file_hash,
    )


# ─── Plain text (MD, TXT, CSV, TSV, RTF) ────────────────────────────

def ingest_text(filepath: str) -> DocumentContent:
    """
    Read plain text files directly.

    For CSV/TSV, also parses into table structure.
    """
    file_hash = compute_file_hash(filepath)
    filename = os.path.basename(filepath)
    ext = os.path.splitext(filepath)[1].lower().lstrip('.')

    # Read the file
    text = ''
    for encoding in ('utf-8', 'utf-8-sig', 'latin-1', 'cp1252'):
        try:
            with open(filepath, 'r', encoding=encoding) as f:
                text = f.read()
            break
        except (UnicodeDecodeError, UnicodeError):
            continue

    if not text:
        # Binary fallback
        with open(filepath, 'rb') as f:
            text = f.read().decode('utf-8', errors='replace')

    # For CSV/TSV, also parse tables
    tables = []
    if ext in ('csv', 'tsv'):
        import csv
        delimiter = '\t' if ext == 'tsv' else ','
        try:
            rows = []
            for line in text.splitlines():
                reader = csv.reader([line], delimiter=delimiter)
                for row in reader:
                    rows.append(row)
            if len(rows) >= 2:
                tables.append(rows)
        except Exception as e:
            logger.warning(f"CSV parsing failed for {filename}: {e}")

    pages = [PageContent(
        page_number=1,
        text=text,
        tables=tables,
        is_scanned=False,
    )]

    return DocumentContent(
        filepath=filepath,
        filename=filename,
        pages=pages,
        page_count=1,
        is_scanned=False,
        file_hash=file_hash,
    )
