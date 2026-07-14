"""
Midway P4 — Disposition Diligence Report (Word deliverable).

Renders the warehouse into a single review-grade Word document:
  1. The Sale — PSA economics (both agreements: price, earnest money, key dates).
  2. Tenant Lease Abstracts — the core facts extracted per tenant.
  3. Missing Documents — the auto-detected executed-instrument gaps.
  4. REA Prohibited Uses — the Exhibit F schedule (grocery/supermarket flagged).

Machine-generated from certification documents; labeled an automated first draft for
review — NOT legal advice.

Usage:
    python diligence_report.py --db data/midway.db --out Midway_Disposition_Diligence.docx
"""
import argparse
import datetime
import os
import sqlite3

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

INK = RGBColor(0x1A, 0x1A, 0x2E)
ACCENT = RGBColor(0x18, 0x5F, 0xA5)
RED = RGBColor(0xC8, 0x10, 0x2E)
NAVY_FILL = "1F3A5F"
CLOUD_FILL = "EEF3F8"
FLAG_FILL = "FCEAEA"

CORE_ORDER = ["instrument_type", "tenant_entity", "landlord", "premises_address",
              "premises_sf", "lease_date", "term_expiration", "base_rent",
              "renewal_options", "permitted_use", "assignment_status",
              "security_deposit", "notice_address_tenant"]
CORE_LABEL = {
    "instrument_type": "Instrument", "tenant_entity": "Tenant", "landlord": "Landlord",
    "premises_address": "Premises", "premises_sf": "SF", "lease_date": "Lease date",
    "term_expiration": "Expiration", "base_rent": "Base rent",
    "renewal_options": "Renewals", "permitted_use": "Use",
    "assignment_status": "Assignment", "security_deposit": "Deposit",
    "notice_address_tenant": "Notice address"}


def _shade(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), fill)
    tcpr.append(shd)


def _title(doc):
    def line(txt, size, color, bold=True, after=2):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(after)
        r = p.add_run(txt); r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = color
        return p
    line("MIDWAY MARKETPLACE", 22, INK)
    line("Disposition Diligence Report — St. Paul, MN", 12, ACCENT, bold=False, after=10)
    line("Partial sale to GSC RE Holdings (H Mart affiliate) · closed 6/18/2026", 10, ACCENT, bold=False)
    banner = line("AUTOMATED FIRST DRAFT FROM CERTIFICATION DOCUMENTS — FOR REVIEW, NOT LEGAL ADVICE",
                  9, RED, after=2)
    banner.paragraph_format.space_before = Pt(8)


def _kv_table(doc, rows, label_w=1.6, val_w=4.9):
    t = doc.add_table(rows=0, cols=2); t.style = "Table Grid"
    for k, v in rows:
        c = t.add_row().cells
        c[0].width = Inches(label_w); c[1].width = Inches(val_w)
        _shade(c[0], CLOUD_FILL)
        r0 = c[0].paragraphs[0].add_run(k); r0.bold = True; r0.font.size = Pt(9)
        c[1].paragraphs[0].add_run(str(v) if v is not None else "—").font.size = Pt(9)
    return t


def build(db_path, out_path):
    con = sqlite3.connect(db_path); con.row_factory = sqlite3.Row
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)
    _title(doc)

    # 1. The Sale (PSA economics)
    doc.add_heading("1 · The Sale — Purchase & Sale Agreements", level=1)
    agrs = con.execute("SELECT * FROM agreement ORDER BY agreement_id").fetchall()
    if not agrs:
        doc.add_paragraph("No PSA extracted.").runs[0].italic = True
    for a in agrs:
        h = doc.add_heading(a["label"] or a["source_file"] or "Agreement", level=2)
        for run in h.runs:
            run.font.color.rgb = INK
        fins = con.execute("SELECT item, value, notes FROM financial_term WHERE agreement_id=?",
                           (a["agreement_id"],)).fetchall()
        dates = con.execute("SELECT item, trigger, duration FROM key_date_deadline WHERE agreement_id=?",
                            (a["agreement_id"],)).fetchall()
        rows = [(f["item"], f["value"] + (f"  ({f['notes']})" if f["notes"] else "")) for f in fins]
        rows += [(d["item"], (d["duration"] or "") + (f" — {d['trigger']}" if d["trigger"] else ""))
                 for d in dates]
        if rows:
            _kv_table(doc, rows)

    # 2. Tenant lease abstracts
    doc.add_heading("2 · Tenant Lease Abstracts", level=1)
    tenants = con.execute("SELECT tenant_id, name FROM lease_tenant ORDER BY name").fetchall()
    for t in tenants:
        facts = {r["field"]: r["value"] for r in con.execute(
            "SELECT field, value FROM lease_abstract WHERE tenant_id=? AND source_page LIKE '[%'",
            (t["tenant_id"],))}
        if not facts:
            continue
        h = doc.add_heading(t["name"], level=2)
        for run in h.runs:
            run.font.color.rgb = INK
        rows = [(CORE_LABEL.get(f, f), facts[f]) for f in CORE_ORDER if f in facts]
        _kv_table(doc, rows)

    # 3. Missing documents
    doc.add_heading("3 · Missing Documents (executed-instrument gaps)", level=1)
    md = con.execute("SELECT item, priority, why_needed FROM missing_document "
                     "WHERE source_or_where LIKE 'auto:%' ORDER BY priority").fetchall()
    if md:
        t = doc.add_table(rows=1, cols=2); t.style = "Table Grid"
        hdr = t.rows[0].cells
        for i, htxt in enumerate(["Missing item", "Priority"]):
            _shade(hdr[i], NAVY_FILL)
            r = hdr[i].paragraphs[0].add_run(htxt); r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.font.size = Pt(9)
        for m in md:
            c = t.add_row().cells
            c[0].width = Inches(5.4); c[1].width = Inches(1.1)
            c[0].paragraphs[0].add_run(m["item"]).font.size = Pt(9)
            c[1].paragraphs[0].add_run(m["priority"] or "").font.size = Pt(9)
    else:
        doc.add_paragraph("No gaps recorded (run missing_docs.py).").runs[0].italic = True

    # 4. REA prohibited uses
    doc.add_heading("4 · REA Prohibited Uses (Exhibit F)", level=1)
    pu = con.execute("SELECT item_no, prohibited_use, exceptions FROM rea_prohibited_use").fetchall()
    if pu:
        note = doc.add_paragraph()
        nr = note.add_run("Diligence flag: any grocery/supermarket prohibition below directly "
                          "affects the buyer's Asian-grocery + food-hall plan — verify anchor-box "
                          "carve-outs in the full REA.")
        nr.italic = True; nr.font.size = Pt(9); nr.font.color.rgb = RED
        t = doc.add_table(rows=1, cols=3); t.style = "Table Grid"
        hdr = t.rows[0].cells
        for i, htxt in enumerate(["#", "Prohibited use", "Exceptions"]):
            _shade(hdr[i], NAVY_FILL)
            r = hdr[i].paragraphs[0].add_run(htxt); r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.font.size = Pt(9)
        for p in pu:
            c = t.add_row().cells
            c[0].width = Inches(0.4); c[1].width = Inches(3.9); c[2].width = Inches(2.2)
            flag = any(w in (p["prohibited_use"] or "").lower()
                       for w in ("grocery", "supermarket", "food market"))
            c[0].paragraphs[0].add_run(str(p["item_no"] or "")).font.size = Pt(9)
            r1 = c[1].paragraphs[0].add_run(p["prohibited_use"] or ""); r1.font.size = Pt(9)
            c[2].paragraphs[0].add_run(p["exceptions"] or "").font.size = Pt(9)
            if flag:
                r1.bold = True; r1.font.color.rgb = RED
                for cell in c:
                    _shade(cell, FLAG_FILL)
    else:
        doc.add_paragraph("No prohibited uses extracted (run extract_rea.py).").runs[0].italic = True

    con.close()
    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--db", default=os.path.join(os.path.dirname(os.path.abspath(__file__)),
                                                 "data", "midway.db"))
    ap.add_argument("--out", default="Midway_Disposition_Diligence.docx")
    args = ap.parse_args()
    p = build(args.db, args.out)
    print(f"Wrote {p}")
