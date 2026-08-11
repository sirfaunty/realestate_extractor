"""
Refinance Diligence Package builder (.docx) — the KA Loan Overview
Platform's end product for one property: loan facilities with balances and
balloons, loan-document provisions by lender category, lease rollover
crossed against loan maturity, lender-relevant lease provisions, and the
open-items register. Verified sources only; every item carries its citation.
"""

import datetime
import os
import re

from docx import Document
from docx.shared import Inches, Pt, RGBColor

from .compendium import (KA_NAVY, KA_RED, KA_BLUE, GREY, NAVY_FILL,
                         CLOUD_FILL, RED_FILL, _shade, _footer, _center_line,
                         _red_rule, _kv_table, _fmt_money, _fmt_sf, _clean)

CATEGORY_LABELS = {
    'transfer_restriction': 'Transfer Restrictions',
    'lender_consent': 'Lender Consent Rights',
    'leasing_approval': 'Leasing Approval Requirements',
    'prepayment': 'Prepayment',
    'financial_covenant': 'Financial Covenants',
    'guaranty_recourse': 'Guaranty / Recourse',
    'change_of_control': 'Change of Control',
    'reporting_requirement': 'Reporting Requirements',
    'reserve_escrow': 'Reserves / Escrows',
    'insurance_condemnation': 'Insurance / Condemnation',
    'future_funding': 'Future Funding',
    'management_agreement': 'Management Agreement',
}


def _table(doc, headers, widths):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        c.width = Inches(widths[i])
        _shade(c, NAVY_FILL)
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    return tbl


def _row(tbl, vals, widths, size=8.5, red_idx=None):
    cells = tbl.add_row().cells
    for i, v in enumerate(vals):
        cells[i].width = Inches(widths[i])
        r = cells[i].paragraphs[0].add_run(_clean(str(v if v not in
                                                     (None, '') else '—')))
        r.font.size = Pt(size)
        if red_idx is not None and i == red_idx and v not in (None, '', '—'):
            r.font.color.rgb = KA_RED
            r.bold = True
    return cells


def build_refi_package(data, out_path):
    prop = data['property']
    pname = prop.get('property_name') or prop.get('property_key')
    today = datetime.date.today().isoformat()

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # ── Title ──
    _center_line(doc, "KRAUS-ANDERSON", 22, KA_NAVY)
    sub = f"Loan Overview Platform  |  {pname}"
    if prop.get('entity_code'):
        sub += f"  ·  Entity {prop['entity_code']}"
    _center_line(doc, sub, 11, KA_BLUE, bold=False, after=10)
    _center_line(doc, "Refinance Diligence Package", 18, KA_NAVY, after=2)
    _center_line(doc, f"Generated {today} from the verified KA portfolio "
                      f"warehouse (loan + lease layers, page-cited)",
                 9, GREY, bold=False, after=2)
    _red_rule(doc)
    banner = _center_line(
        doc, "COMPILED FROM VERIFIED, PAGE-CITED EXTRACTIONS — "
             "VERIFY AGAINST SOURCE INSTRUMENTS BEFORE LEGAL RELIANCE",
        9, KA_RED, after=2)
    banner.paragraph_format.space_before = Pt(8)

    # ── Executive snapshot ──
    doc.add_heading("Executive Snapshot", level=1)
    cur = [f for f in data['facilities']
           if (f.get('loan_status') or '').lower() == 'current']
    snap_pairs = [
        ("Current facilities", f"{len(cur)} of {len(data['facilities'])} total"),
        ("Earliest current maturity", data.get('earliest_current_maturity')),
        ("Occupied tenancies", data['occupied_count']),
        ("Leases expiring before maturity",
         f"{len(data['rollover_before_maturity'])}"
         + (' — see Rollover section' if data['rollover_before_maturity'] else
            ' (all occupied leases run past maturity)')),
        ("Open items (loan workstreams)", len(data['open_items'])),
        ("Rent roll snapshot", data.get('rent_roll_snapshot')),
    ]
    _kv_table(doc, [(k, v) for k, v in snap_pairs if v not in (None, '')])

    # ── Facilities ──
    doc.add_heading("Loan Facilities", level=1)
    w = [1.15, 1.35, 1.0, 0.6, 0.85, 0.75, 0.8]
    tbl = _table(doc, ("Facility", "Lender", "Principal", "Rate",
                       "Maturity", "Status", "Balance"), w)
    for f in data['facilities']:
        bal = f.get('balance') or {}
        _row(tbl, (f['facility_id'],
                   f.get('lender'),
                   _fmt_money(f.get('original_principal')),
                   f"{f['interest_rate_pct']}%" if f.get('interest_rate_pct') else None,
                   f.get('maturity_date'),
                   f.get('loan_status'),
                   _fmt_money(bal.get('balance'))
                   and f"{_fmt_money(bal.get('balance'))} ({bal.get('asof_date')})"),
             w)
    for f in data['facilities']:
        notes = []
        if f.get('balloon') and f['balloon'].get('balloon_balance'):
            notes.append(f"balloon {_fmt_money(f['balloon']['balloon_balance'])} "
                         f"at {f['balloon'].get('maturity_date')}")
        if f.get('is_cross_collateralized'):
            notes.append(f"cross-collateralized "
                         f"({f.get('collateral_property_count')} properties)")
        if f.get('principal_confidence') and 'verified' not in f['principal_confidence']:
            notes.append(f"principal confidence: {f['principal_confidence']}")
        if f.get('rate_confidence') and f['rate_confidence'] not in ('verified',):
            notes.append(f"rate confidence: {f['rate_confidence']}")
        if notes:
            p = doc.add_paragraph()
            r = p.add_run(_clean(f"{f['facility_id']}: " + '; '.join(notes)))
            r.font.size = Pt(8.5)
            r.font.color.rgb = GREY

    # ── Loan document provisions ──
    if data['loan_provisions']:
        doc.add_heading("Loan Document Provisions (by lender category)", level=1)
        for cat in sorted(data['loan_provisions']):
            label = CATEGORY_LABELS.get(cat, cat.replace('_', ' ').title())
            h = doc.add_heading(f"{label} ({len(data['loan_provisions'][cat])})",
                                level=2)
            for run in h.runs:
                run.font.color.rgb = KA_NAVY
            for p_ in data['loan_provisions'][cat]:
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.15)
                if p_.get('why_it_matters'):
                    r = para.add_run(_clean(p_['why_it_matters'].strip()) + '  ')
                    r.font.size = Pt(9.5)
                if p_.get('evidence'):
                    ev = p_['evidence'].strip()
                    if len(ev) > 400:
                        ev = ev[:400] + ' …'
                    r = para.add_run(_clean(f'“{ev}”'))
                    r.font.size = Pt(9)
                    r.italic = True
                cite = doc.add_paragraph()
                cite.paragraph_format.left_indent = Inches(0.15)
                cite.paragraph_format.space_after = Pt(6)
                r = cite.add_run(_clean(
                    f"Source: {p_.get('source_file') or '(loan layer)'}"
                    + (f" · {p_.get('kind')}" if p_.get('kind') else '')))
                r.font.size = Pt(8)
                r.font.color.rgb = GREY

    # ── Lease rollover vs maturity ──
    doc.add_heading("Lease Rollover vs Loan Maturity", level=1)
    note = doc.add_paragraph()
    note.add_run(
        f"Earliest current-facility maturity: "
        f"{data.get('earliest_current_maturity') or 'n/a'}. Occupied leases "
        f"expiring on or before that date are flagged."
    ).font.size = Pt(9.5)
    w2 = [1.9, 0.8, 0.9, 1.1, 1.4]
    tbl2 = _table(doc, ("Tenant", "SF", "Status", "Expiration",
                        "Expires before maturity"), w2)
    for t in data['tenancies']:
        if (t.get('status') or '').lower() not in ('current', 'active',
                                                   'occupied'):
            continue
        _row(tbl2, (t['trade_name'] or t['tenant_key'],
                    _fmt_sf(t.get('sf')),
                    (t.get('status') or '').title(),
                    t.get('expiration'),
                    'YES — rollover risk' if t['expires_before_maturity'] else 'no'),
             w2, red_idx=4 if t['expires_before_maturity'] else None)

    # ── Lender-relevant lease provisions ──
    doc.add_heading("Lender-Relevant Lease Provisions", level=1)
    exp = doc.add_paragraph()
    exp.add_run(
        "Provisions flagged REFI IMPACT by the aggregator appear first; the "
        "remainder are selected by category (SNDA/subordination, assignment, "
        "termination rights, co-tenancy, exclusives, purchase options, "
        "renewals, guaranties) — each labeled with which net caught it."
    ).font.size = Pt(9.5)
    for t in data['tenancies']:
        provs = t['refi_provisions'] + t['lender_provisions']
        if not provs:
            continue
        h = doc.add_heading(_clean(t['trade_name'] or t['tenant_key'])
                            + f"  ({len(provs)})", level=2)
        for run in h.runs:
            run.font.color.rgb = KA_NAVY
        w3 = [2.1, 3.0, 1.4]
        tbl3 = _table(doc, ("Provision", "Note", "Source"), w3)
        for p_ in t['refi_provisions']:
            _row(tbl3, (p_['category'],
                        f"REFI IMPACT: {p_['refi_impact'].strip()}",
                        (p_.get('source') or '')
                        + (f" pp.{p_['source_pages']}" if p_.get('source_pages') else '')),
                 w3, red_idx=1)
        for p_ in t['lender_provisions']:
            d = (p_.get('detail') or '').strip()
            if len(d) > 220:
                d = d[:220] + ' …'
            _row(tbl3, (p_['category'], d,
                        (p_.get('source') or '')
                        + (f" pp.{p_['source_pages']}" if p_.get('source_pages') else '')),
                 w3)

    # ── Open items ──
    if data['open_items']:
        doc.add_heading("Open Items (loan workstreams)", level=1)
        w4 = [1.2, 2.6, 2.0, 0.7]
        tbl4 = _table(doc, ("Workstream", "Item", "Why", "Priority"), w4)
        for oi in data['open_items']:
            _row(tbl4, (oi.get('workstream'), oi.get('item'), oi.get('why'),
                        oi.get('priority')), w4,
                 red_idx=3 if str(oi.get('priority', '')).lower()
                 in ('high', '1', 'p1') else None)

    _footer(doc.sections[0],
            f"{pname} Refinance Diligence Package · verified-source render")
    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)
    return {'out': out_path,
            'facilities': len(data['facilities']),
            'loan_provisions': sum(len(v) for v in data['loan_provisions'].values()),
            'lease_provisions': sum(len(t['refi_provisions'])
                                    + len(t['lender_provisions'])
                                    for t in data['tenancies'])}


def refi_filename(property_name):
    safe = re.sub(r'[^A-Za-z0-9]+', '_', property_name or 'Property').strip('_')
    return (f"{safe}_Refinance_Diligence_Package_"
            f"{datetime.date.today().isoformat()}.docx")
